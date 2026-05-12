# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

IMG_DIR = r"C:\Users\Duy\Desktop\dvfinal\docs\tableau"

# ── helpers ──────────────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

def h2(text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

def h3(text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(11)

def body(text):
    doc.add_paragraph(text)

def bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

def blank():
    doc.add_paragraph()

def insert_image(filename, caption=""):
    path = os.path.join(IMG_DIR, filename)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(6.0))
        if caption:
            p = doc.add_paragraph(caption)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].italic = True
    else:
        doc.add_paragraph(f"[Image not found: {filename}]")

def idiom_table(rows):
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.style = 'Table Grid'
    for i, (lbl, content) in enumerate(rows):
        cell_l = tbl.cell(i, 0)
        cell_r = tbl.cell(i, 1)
        r = cell_l.paragraphs[0].add_run(lbl); r.bold = True
        cell_r.paragraphs[0].add_run(content)

def eval_section(express_bullets, effect_bullets, conclude_express, conclude_effect, phanTich_bullets):
    h3("B. Đánh giá biểu đồ")
    p = doc.add_paragraph(); p.add_run("1. Tính biểu đạt (Expressiveness):").bold = True
    for b in express_bullets:
        bullet(b)
    body(conclude_express)
    p = doc.add_paragraph(); p.add_run("2. Tính hiệu quả (Effectiveness):").bold = True
    for b in effect_bullets:
        bullet(b)
    if conclude_effect:
        body(conclude_effect)
    h3("C. Phân tích biểu đồ")
    for b in phanTich_bullets:
        bullet(b)

# ══════════════════════════════════════════════════════════════════════════════
# TASK 5
# ══════════════════════════════════════════════════════════════════════════════
h1("Domain task 5")
body("Câu hỏi: Giá niêm yết phân bổ như thế nào theo từng borough? Loại phòng ảnh hưởng thế nào đến phân phối giá?")
blank()

h2("5.1 – Box-and-Whisker Plot: Phân phối giá niêm yết theo borough")
idiom_table([
    ("Idiom", "Box-and-Whisker Plot"),
    ("What",
     "Borough (neighbourhood_group_cleansed): Categorical\n"
     "Price: Quantitative\n"
     "Room Type: Categorical"),
    ("How\nEncode",
     "Mark: Line (box plot – whisker, hộp IQR, đường median)\n"
     "Channel:\n"
     "  Pos X: Borough (Categorical – phân biệt 5 borough)\n"
     "  Pos Y: Price (Quantitative – IQR, median, whiskers)\n"
     "  Hue Color: Room Type (Categorical – 4 loại phòng)"),
    ("How\nManipulate", "Selection: hover để xem chi tiết (median, Q1, Q3, min, max, count)"),
    ("How\nFacet", "—"),
    ("How\nReduce", "Filter: price_is_outlier = False (loại extreme outliers bằng IQR method)"),
    ("Why",
     "compare → discover\n"
     "So sánh phân phối giá đầy đủ giữa 5 borough. Box plot là idiom duy nhất hiển thị Q1, median, Q3, whiskers đồng thời — bar chart chỉ cho median, histogram chỉ cho 1 nhóm."),
    ("Scale", "Main key: 5 (borough)\nColor key: 4 (room type)\nItems: ~21,000 listing (sau lọc)"),
])
blank()
insert_image("Task 5.1 - Price Distribution by Borough.png",
             "Hình 5.1. Box plot phân phối giá niêm yết theo borough tại NYC")
blank()
eval_section(
    express_bullets=[
        "PosX: phân biệt 5 borough (C) → phù hợp. Spatial region là channel hiệu quả nhất cho categorical attribute.",
        "PosY: thể hiện phạm vi phân phối giá (Q1, Median, Q3, whiskers) → phù hợp Quantitative. Position on common scale là channel chính xác nhất cho quantitative.",
        "Hue Color: phân biệt loại phòng (C) → phù hợp. Số màu = 4 ≤ 7 → nằm trong giới hạn discriminability của HUE channel.",
    ],
    conclude_express="Kết luận: Các channel đúng với bản chất dữ liệu. Box plot là idiom chuẩn mực nhất để biểu diễn phân phối liên tục theo nhóm — đặc biệt khi cần so sánh IQR chứ không chỉ median.",
    effect_bullets=[
        "Accuracy: PosY (Position on common scale) là channel chính xác nhất theo channel effectiveness ranking — xếp hạng 1 trong hierarchy. Hue Color dùng cho categorical không liên quan đến sai số định lượng.",
        "Discriminability: 4 màu cho 4 room type trong phạm vi ≤ 7 → mắt người phân biệt hoàn toàn tốt. Whiskers và box IQR hiển thị rõ ràng trên nền trắng.",
        "Separability: PosX, PosY và Color tách biệt hoàn toàn. Filter price_is_outlier = False loại bỏ extreme outliers giúp chart dễ đọc hơn.",
    ],
    conclude_effect="",
    phanTich_bullets=[
        "Manhattan có median price cao nhất (~$150–200) và IQR rộng nhất — tồn tại cả phân khúc bình dân lẫn cao cấp trong cùng một borough.",
        "Bronx và Staten Island có IQR hẹp và median thấp (~$92–99) — thị trường giá ổn định, phù hợp du khách cần ngân sách dự đoán được.",
        "Loại phòng Entire home/apt luôn có box cao hơn Private room ở mọi borough — sự phân tầng giá theo loại phòng nhất quán trên toàn thành phố.",
        "Khuyến nghị: Du khách ngân sách trung bình nên cân nhắc Brooklyn — vị trí gần Manhattan nhưng median giá thấp hơn đáng kể.",
    ]
)
blank()

h2("5.2 – Bar Chart: Median giá niêm yết theo borough")
idiom_table([
    ("Idiom", "Bar Chart"),
    ("What", "Borough: Categorical\nMedian Price: Quantitative (aggregate: MEDIAN)"),
    ("How\nEncode",
     "Mark: Bar\n"
     "Channel:\n"
     "  Pos X: Borough (Categorical)\n"
     "  Pos Y: MEDIAN(Price) (Quantitative – length từ gốc 0)\n"
     "  Hue Color: Borough (Categorical – redundant encoding)\n"
     "  Label: giá trị MEDIAN trên đầu cột\n"
     "  Order: sort descending theo median price"),
    ("How\nManipulate", "Hover để xem median price chính xác theo từng borough"),
    ("How\nFacet", "—"),
    ("How\nReduce", "Filter: price_is_outlier = False\nAggregate: MEDIAN(price) theo borough"),
    ("Why",
     "compare → summarize\n"
     "Tóm tắt median giá thành 1 số đại diện dễ so sánh. Sort descending giúp nhận diện borough đắt/rẻ nhất ngay lập tức. Bar chart phù hợp khi task là Compare giá trị tuyệt đối."),
    ("Scale", "Main key: 5 (borough)\nItems: ~21,000 listing (sau lọc)"),
])
blank()
insert_image("Task 5.2 - Median Price by Borough.png",
             "Hình 5.2. Bar chart median giá niêm yết theo borough tại NYC")
blank()
eval_section(
    express_bullets=[
        "PosX: phân biệt borough (C) → đúng.",
        "PosY (Length từ gốc 0): thể hiện MEDIAN Price (Q) → đúng. Baseline = 0 bắt buộc để length channel có ý nghĩa chính xác.",
        "Hue Color: phân biệt borough (C) → đúng. Redundant encoding (cả PosX lẫn Color encode Borough) gửi thông điệp mạnh hơn.",
    ],
    conclude_express="Kết luận: Mọi channel phù hợp và đúng với bản chất dữ liệu.",
    effect_bullets=[
        "Accuracy: Kênh Length xếp hạng thứ 3 trong channel effectiveness ranking (sau Position on common scale và Position on unaligned scale). Label số ở đầu cột loại bỏ hoàn toàn sai số cảm nhận thị giác.",
        "Discriminability: 5 cột màu khác nhau, phân cách rõ ràng, sort descending giúp so sánh nhanh từ đắt đến rẻ.",
        "Separability: PosX, PosY và Color tách biệt hoàn toàn.",
    ],
    conclude_effect="",
    phanTich_bullets=[
        "Manhattan ($200) đắt hơn gấp đôi so với Bronx ($93) — mức chênh lệch rất lớn, cho thấy sự phân cực giá rõ rệt theo vị trí địa lý tại NYC.",
        "Brooklyn ($129) đứng thứ hai — lựa chọn cân bằng giữa vị trí và chi phí, rẻ hơn Manhattan $71/đêm.",
        "Queens ($100), Staten Island ($99) và Bronx ($93) có median tương đương nhau — phù hợp cho du khách ngân sách thấp.",
        "Khuyến nghị: Với nhóm du khách muốn cân bằng tiện lợi và tiết kiệm, Brooklyn là lựa chọn tối ưu; với ngân sách tối thiểu, Queens hoặc Bronx phù hợp nhất.",
    ]
)
blank()

# ══════════════════════════════════════════════════════════════════════════════
# TASK 6
# ══════════════════════════════════════════════════════════════════════════════
h1("Domain task 6")
body("Câu hỏi: Listing nào có giá bất thường (outlier)? Listing nào có giá thấp nhưng đánh giá cao (good deal)?")
blank()

h2("6.1 – Point Map: Phân bố listing theo giá bất thường")
idiom_table([
    ("Idiom", "Point Map (Geographic Scatter Map)"),
    ("What",
     "Longitude: Quantitative/Geographic (Key)\n"
     "Latitude: Quantitative/Geographic (Key)\n"
     "price_is_outlier (derived): Categorical (True/False)\n"
     "Price: Quantitative\n"
     "Room Type: Categorical"),
    ("How\nEncode",
     "Mark: Point (Circle)\n"
     "Channel:\n"
     "  Pos X: Longitude (Geographic)\n"
     "  Pos Y: Latitude (Geographic)\n"
     "  Hue Color: price_is_outlier (Categorical – đỏ=True, xanh=False)\n"
     "  Size: Price (Quantitative – điểm lớn = giá cao)"),
    ("How\nManipulate", "Hover để xem id, price, room_type, neighbourhood, price_is_outlier"),
    ("How\nFacet", "—"),
    ("How\nReduce",
     "Giải pháp overplotting: giảm opacity điểm False (xanh) để outlier (đỏ) nổi bật;\n"
     "hoặc filter theo borough để zoom vào từng khu vực"),
    ("Why",
     "locate → explore\n"
     "Point Map là idiom duy nhất trả lời câu hỏi 'ở đâu' — không thể thay thế bằng bar chart hay box plot cho phân tích spatial distribution."),
    ("Scale", "Main key: ~21,000 listings\nColor: 2 (True/False outlier)"),
])
blank()
insert_image("Task 6.1 - Price Outlier Map.png",
             "Hình 6.1. Bản đồ phân bố listing theo giá bất thường tại NYC")
blank()
eval_section(
    express_bullets=[
        "PosX (Longitude), PosY (Latitude): vị trí địa lý thực tế → phù hợp Quantitative/Geographic. Geographic position là cách duy nhất đúng để encode tọa độ không gian.",
        "Hue Color: phân biệt nhị phân True/False outlier (C) → phù hợp. Chỉ 2 màu → discriminability tuyệt đối.",
        "Size: thể hiện mức giá (Q) → phù hợp — điểm lớn hơn = giá cao hơn, nhất quán với convention.",
    ],
    conclude_express="Kết luận: Point Map là idiom tối ưu để phân tích dữ liệu có thành phần địa lý.",
    effect_bullets=[
        "Accuracy: PosX, PosY (geographic position) có độ chính xác tuyệt đối về không gian. Kênh Size (Area) xếp thấp trong channel effectiveness ranking — khó ước lượng chênh lệch giá chính xác. Trade-off chấp nhận được vì mục tiêu là Locate, không phải Compare.",
        "Discriminability: 2 màu nhị phân (đỏ/xanh) rất dễ phân biệt. Với ~21k điểm, các vùng dày đặc bị overlap nghiêm trọng. Giải pháp: giảm opacity điểm False (xanh), hoặc áp dụng filter theo borough.",
        "Separability: PosX/PosY và Color tách biệt tốt. Size và Color có thể tương tác nhẹ khi điểm lớn che điểm nhỏ.",
    ],
    conclude_effect="",
    phanTich_bullets=[
        "Outlier giá cao (điểm đỏ lớn) tập trung dày đặc ở Manhattan — đặc biệt Midtown và Upper East Side, khẳng định Manhattan là thị trường giá cao và biến động nhất.",
        "Bronx và Staten Island hầu như không có điểm đỏ — thị trường giá bình ổn, an toàn hơn cho du khách về khả năng dự đoán giá.",
        "Brooklyn có một số outlier tập trung ở Brooklyn Heights, DUMBO — nơi host định giá cao hơn mức phổ thông do view đẹp và gần Manhattan.",
        "Khuyến nghị: Du khách nên kết hợp với Task 5 để chọn borough; tránh listing Manhattan không có nhiều reviews và có giá bất thường cao.",
    ]
)
blank()

h2("6.2 – Scatter Plot: Giá vs Rating — Phát hiện Good Deal")
idiom_table([
    ("Idiom", "Scatter Plot"),
    ("What",
     "Price: Quantitative (Key trục X)\n"
     "review_scores_rating: Quantitative (Key trục Y)\n"
     "good_deal_flag (derived): Categorical (Good Deal / Normal)\n"
     "number_of_reviews: Quantitative (Size – proxy độ tin cậy)"),
    ("How\nEncode",
     "Mark: Point (Circle)\n"
     "Channel:\n"
     "  Pos X: Price (Quantitative)\n"
     "  Pos Y: review_scores_rating (Quantitative)\n"
     "  Hue Color: good_deal_flag (Categorical – xanh=Good Deal, xám=Normal)\n"
     "  Size: number_of_reviews (Quantitative)\n"
     "Reference Lines:\n"
     "  Dọc (trục X): MEDIAN(price) – phân cách giá cao/thấp\n"
     "  Ngang (trục Y): Constant = 4.8 – ngưỡng Good Deal Flag"),
    ("How\nManipulate", "Hover để xem giá, rating, room_type, neighbourhood, number_of_reviews"),
    ("How\nFacet", "Superimpose: Reference Lines đặt lên scatter plot"),
    ("How\nReduce", "Filter: number_of_reviews ≥ 5 (loại listing ít đánh giá, không đủ tin cậy)"),
    ("Why",
     "discover → compare\n"
     "Scatter plot tối ưu cho phân tích 2Q correlation. 4 góc phần tư từ reference lines giúp nhận diện Good Deal (góc trên trái: price < median VÀ rating ≥ 4.8)."),
    ("Scale", "Color: 2 (Good Deal / Normal)\nItems: listing với ≥ 5 reviews, price_is_outlier = False"),
])
blank()
insert_image("Task 6.2 - Price vs Rating.png",
             "Hình 6.2. Scatter plot giá niêm yết vs. điểm đánh giá — phát hiện Good Deal")
blank()
eval_section(
    express_bullets=[
        "PosX (Price – Q): vị trí ngang thể hiện giá → đúng.",
        "PosY (Rating – Q): vị trí dọc thể hiện chất lượng → đúng. Rating cao hơn = vị trí cao hơn nhất quán với convention 'up = good'.",
        "Hue Color: phân biệt Good Deal / Normal (C) → đúng. Chỉ 2 giá trị, discriminability tuyệt đối.",
        "Size: thể hiện Number of Reviews (Q) → đúng. Size phù hợp cho Q thứ cấp (proxy trust level).",
    ],
    conclude_express="Kết luận: Scatter plot là idiom tối ưu để phân tích mối quan hệ giữa 2 biến định lượng với phân nhóm categorical.",
    effect_bullets=[
        "Accuracy: PosX và PosY (position on common scale) — channel chính xác nhất. Kênh Size (Area) — accuracy thấp hơn, khó so sánh chính xác số reviews; đây là channel phụ nên chấp nhận được.",
        "Discriminability: 2 màu (xanh/xám) rất dễ phân biệt. Reference lines tạo 4 góc phần tư rõ ràng, giúp người xem định vị ngay vùng 'good deal' (góc trên-trái). Overplotting tại rating 4.8–5.0 là hạn chế — có thể cải thiện bằng jitter hoặc transparency.",
        "Separability: PosX, PosY, Color và Size tách biệt tốt.",
    ],
    conclude_effect="",
    phanTich_bullets=[
        "Vùng góc trên-trái (price < median, rating ≥ 4.8) là vùng 'Good Deal' — tập trung nhiều điểm xanh, chủ yếu thuộc Brooklyn và Queens.",
        "Manhattan có nhiều listing ở phần bên phải chart (giá cao) với rating không tương xứng — không phải lựa chọn tối ưu về cost-efficiency.",
        "Listing 'Good Deal' có Size lớn (nhiều reviews) là những nơi đã được kiểm chứng bởi nhiều khách — độ tin cậy cao.",
        "Quan trọng: Không có tương quan dương rõ ràng giữa giá cao và rating cao — nhiều listing Brooklyn/Queens giá thấp vẫn đạt rating 4.7–5.0, chứng minh giá không phải là proxy của chất lượng trên Airbnb.",
        "Khuyến nghị: Ưu tiên listing trong vùng 'Good Deal' tại Brooklyn/Queens với number_of_reviews ≥ 20–30.",
    ]
)
blank()

# ══════════════════════════════════════════════════════════════════════════════
# TASK 7
# ══════════════════════════════════════════════════════════════════════════════
h1("Domain task 7")
body("Câu hỏi: Borough và loại phòng nào có chi phí mỗi người thấp nhất? Sức chứa tác động thế nào đến giá/người?")
blank()

h2("7.1 – Grouped Bar Chart: Median chi phí mỗi người theo borough và loại phòng")
idiom_table([
    ("Idiom", "Grouped Bar Chart"),
    ("What",
     "Borough: Categorical (Key nhóm chính)\n"
     "Room Type: Categorical (Key nhóm phụ)\n"
     "price_per_person (derived): Quantitative (price / accommodates)"),
    ("How\nEncode",
     "Mark: Bar (grouped – Stack Marks OFF)\n"
     "Channel:\n"
     "  Pos X: Borough × Room Type (Categorical – tạo cụm cột cạnh nhau)\n"
     "  Pos Y: MEDIAN(price_per_person) (Quantitative – length)\n"
     "  Hue Color: Room Type (Categorical – 4 màu)\n"
     "  Label: giá trị median ở đầu mỗi cột\n"
     "Note: Analysis → Stack Marks → Off (bắt buộc để tạo Grouped, không phải Stacked)"),
    ("How\nManipulate", "Hover để xem median price/person theo từng borough × room type"),
    ("How\nFacet", "—"),
    ("How\nReduce", "Filter: price_is_outlier = False\nAggregate: MEDIAN(price_per_person) theo borough × room type"),
    ("Why",
     "compare → summarize\n"
     "So sánh chi phí/người theo 2 chiều categorical đồng thời (borough × room type). Grouped bar (không phải Stacked) vì task là Compare giá trị tuyệt đối, không phải part-to-whole."),
    ("Scale", "Main key: 5 borough × 4 room type = 20 cột\nItems: ~21,000 listing (sau lọc)"),
])
blank()
insert_image("Task 7.1 - Price per Person by Borough.png",
             "Hình 7.1. Grouped bar chart median giá/người theo borough và loại phòng")
blank()
eval_section(
    express_bullets=[
        "PosX: phân biệt borough (C) → đúng.",
        "PosY (Length từ 0): thể hiện MEDIAN price/person (Q) → đúng. Length là channel chính xác nhất cho Q.",
        "Hue Color: phân biệt room type (C) trong grouped bar → đúng.",
    ],
    conclude_express="Kết luận: Grouped Bar Chart là lựa chọn đúng đắn khi cần so sánh đồng thời theo 2 chiều categorical.",
    effect_bullets=[
        "Accuracy: Length (từ gốc 0) — channel có độ chính xác cao. Label số trên đầu cột loại bỏ sai số cảm nhận hoàn toàn.",
        "Discriminability: 4 màu trong phạm vi ≤ 7 → phân biệt tốt. Grouped layout rõ ràng hơn stacked bar.",
        "Separability: PosX, PosY và Color hoạt động độc lập tốt.",
    ],
    conclude_effect="",
    phanTich_bullets=[
        "Entire home/apt có price/person cao nhất nhìn chung — nhưng khi chia theo nhiều người (4–6 người), đây vẫn cạnh tranh với Private room ở Manhattan.",
        "Bronx và Staten Island có price/person thấp nhất trong mọi room type — lựa chọn tối ưu cho nhóm khách ngân sách thấp.",
        "Private room và Shared room luôn có price/person thấp hơn Entire home — phù hợp cho khách solo hoặc đôi.",
        "Khuyến nghị: Nhóm 4–6 người nên so sánh Entire home/apt ở Queens/Bronx với Private room ở Brooklyn — kết hợp với Task 7.2 để xác định crossover point.",
    ]
)
blank()

h2("7.2 – Scatter Plot: Sức chứa vs Chi phí mỗi người")
idiom_table([
    ("Idiom", "Scatter Plot với Trend Line"),
    ("What",
     "accommodates: Quantitative (Key trục X – sức chứa 1–16)\n"
     "price_per_person (derived): Quantitative (Value trục Y)\n"
     "Room Type: Categorical (Color)"),
    ("How\nEncode",
     "Mark: Point (Circle)\n"
     "Channel:\n"
     "  Pos X: accommodates (Quantitative – giữ từng giá trị riêng 1–16)\n"
     "  Pos Y: MEDIAN(price_per_person) (Quantitative)\n"
     "  Hue Color: Room Type (Categorical – 4 màu)"),
    ("How\nManipulate", "Hover để xem accommodates, median price/person, room_type, count"),
    ("How\nFacet", "Superimpose: Trend Line (Linear) cho từng room type"),
    ("How\nReduce", "Filter: accommodates ≤ 16, price_is_outlier = False"),
    ("Why",
     "discover → explore\n"
     "Phát hiện xu hướng giảm của price/person khi sức chứa tăng (economies of scale). Scatter + Trend Line là idiom chuẩn để phân tích relationship giữa 2 biến Q, đặc biệt để xác định crossover point giữa các room type."),
    ("Scale", "Key: 16 (giá trị accommodates 1–16)\nColor key: 4 room type"),
])
blank()
insert_image("Task 7.2 - Accommodates vs Price per Perso.png",
             "Hình 7.2. Mối quan hệ giữa sức chứa và chi phí mỗi người")
blank()
eval_section(
    express_bullets=[
        "PosX (Accommodates – Q/Ordinal): position thích hợp.",
        "PosY (MEDIAN price/person – Q): position thích hợp.",
        "Hue Color: Room Type (C) → đúng, phân biệt 3–4 loại phòng.",
        "Trend Line: hiển thị xu hướng tổng thể → tăng thêm thông tin về dependency.",
    ],
    conclude_express="Kết luận: Scatter plot với Trend Line là lựa chọn chuẩn mực để phân tích mối quan hệ giữa 2 biến Q và so sánh xu hướng giữa các nhóm.",
    effect_bullets=[
        "Accuracy: PosX và PosY (position on common scale) — rất chính xác. Trend Line thể hiện xu hướng dốc xuống rõ ràng. MEDIAN aggregation giúp giảm nhiễu so với raw data.",
        "Discriminability: 3–4 màu dễ phân biệt. Số lượng điểm vừa phải (~42–64 marks) không bị clutter.",
        "Separability: PosX, PosY và Color tách biệt hoàn toàn. Trend lines không làm rối các điểm dữ liệu thực.",
    ],
    conclude_effect="",
    phanTich_bullets=[
        "Trend line của Entire home/apt dốc xuống mạnh nhất — xác nhận 'economies of scale': càng nhiều người chia phòng, giá mỗi người càng giảm. Sweet spot rõ ràng ở 4–6 người.",
        "Private room gần như phẳng — sức chứa không ảnh hưởng nhiều đến price/person.",
        "Crossover point: Với 4+ người, Entire home trở nên rẻ hơn Private room tính theo đầu người — insight quan trọng nhất của Task 7.",
        "Khuyến nghị: Nhóm ≥ 4 người nên ưu tiên Entire home/apt, ưu tiên Queens/Brooklyn để tối ưu cả vị trí lẫn chi phí mỗi người.",
    ]
)
blank()

# ══════════════════════════════════════════════════════════════════════════════
# TASK 8
# ══════════════════════════════════════════════════════════════════════════════
h1("Domain task 8")
body("Câu hỏi: Tỷ lệ lấp đầy thay đổi như thế nào theo mùa? Chính sách minimum nights ảnh hưởng thế nào?")
blank()

h2("8.1 – Line Chart: Tỷ lệ lấp đầy theo tháng")
idiom_table([
    ("Idiom", "Line Chart"),
    ("What",
     "Month: Ordinal (tháng 1–12, cyclic – seasonality lặp lại theo năm)\n"
     "occupancy_rate_pct (derived): Quantitative (AVG(is_booked) × 100)"),
    ("How\nEncode",
     "Mark: Line + Point\n"
     "Channel:\n"
     "  Pos X: Month (Ordinal, Discrete – tiến trình thời gian tháng 1–12)\n"
     "  Pos Y: Occupancy Rate (%) (Quantitative)\n"
     "Reference Line: Average (ngang – mức trung bình tổng thể ~30%)"),
    ("How\nManipulate", "Hover để xem tháng, occupancy rate (%) cụ thể"),
    ("How\nFacet", "Superimpose: Reference Line (average occupancy ngang)"),
    ("How\nReduce",
     "Filter: Year = 2026 (dữ liệu calendar 2025 chỉ có từ tháng 11, không đủ để phân tích seasonality)"),
    ("Why",
     "discover → browse\n"
     "Line chart (không phải bar) vì Month là Ordinal có thứ tự – đường kết nối nhấn mạnh tính liên tục và xu hướng thời gian, không phải giá trị tại từng điểm rời rạc."),
    ("Scale", "Key: 12 (tháng); Items: dữ liệu 2026 (tháng 1–11)"),
])
blank()
insert_image("Task 8.1 - Monthly Occupancy Rate.png",
             "Hình 8.1. Tỷ lệ lấp đầy (occupancy rate) theo tháng tại NYC (2026)")
blank()
eval_section(
    express_bullets=[
        "PosX (Month – Ordinal): trục thời gian có thứ tự → đúng. Không dùng Categorical vì mất thứ tự thời gian.",
        "PosY (Occupancy Rate – Q): thể hiện mức độ lấp đầy → đúng.",
        "Reference Line Average: cung cấp mốc tham chiếu tổng thể → tăng thêm thông tin phân tích.",
    ],
    conclude_express="Kết luận: Line chart là idiom tối ưu cho dữ liệu chuỗi thời gian, đặc biệt khi cần phát hiện trend và seasonality.",
    effect_bullets=[
        "Accuracy: PosY (position on common scale) — rất chính xác. Người xem dễ dàng so sánh mức occupancy giữa các tháng.",
        "Discriminability: Reference line Average tạo mốc tham chiếu để xác định tháng nào trên/dưới trung bình.",
        "Separability: PosX và PosY tách biệt hoàn toàn.",
    ],
    conclude_effect="",
    phanTich_bullets=[
        "Tháng 1–4 là low season rõ rệt (occupancy ~20–21%) — cơ hội tốt cho khách muốn giá thấp và dễ tìm phòng.",
        "Từ tháng 5 trở đi, occupancy tăng dần và đạt đỉnh vào tháng 11 (~40%) — phản ánh nhu cầu du lịch mùa thu và dịp lễ Thanksgiving.",
        "Dữ liệu sử dụng: calendar 2026 (tháng 1–11); dữ liệu 2025 chỉ có từ tháng 11 nên không đủ để so sánh liên năm.",
        "Khuyến nghị cho du khách: Đặt phòng tháng 1–3 để có giá thấp nhất; tháng 11 cần đặt sớm do cầu rất cao.",
        "Khuyến nghị cho host: Tháng 1–4 nên linh hoạt hóa chính sách (giảm minimum nights, giảm giá nhẹ); tháng 11–12 có thể tăng giá.",
    ]
)
blank()

h2("8.2 – Heatmap: Tỷ lệ lấp đầy theo tháng và chính sách minimum nights")
idiom_table([
    ("Idiom", "Heatmap (Matrix Chart)"),
    ("What",
     "Month: Ordinal (tháng 1–12, cyclic)\n"
     "minimum_nights_group (derived): Categorical (Short ≤3 / Medium 4–7 / Long >7)\n"
     "occupancy_rate_pct (derived): Quantitative"),
    ("How\nEncode",
     "Mark: Square (Cell)\n"
     "Channel:\n"
     "  Pos X: Month (Ordinal – 12 cột)\n"
     "  Pos Y: Minimum Nights Group (Categorical – Short/Medium/Long, top→bottom)\n"
     "  Color (Luminance/Sequential): Occupancy Rate (Quantitative – đậm=cao, nhạt=thấp)"),
    ("How\nManipulate", "Hover để xem month, Minimum Nights Group, occupancy_rate_pct (%)"),
    ("How\nFacet", "—"),
    ("How\nReduce",
     "Binning minimum_nights → 3 nhóm (giảm số chiều từ hàng trăm giá trị → 3 nhóm)\n"
     "Sort thủ công: Short → Medium → Long (thứ tự logic nghiệp vụ, không alphabetical)"),
    ("Why",
     "discover → compare\n"
     "Phát hiện pattern 2 chiều đồng thời (month × minimum_nights_group). Heatmap theo nguyên lý '2 keys → heatmap' — encode 3 biến trong cấu trúc ma trận. Color Luminance/Sequential đúng cho quantitative trong heatmap vì kích thước ô đồng đều loại bỏ bias từ Area."),
    ("Scale", "Key: 12 (tháng) × 3 (nhóm) = 36 ô\nBin: 3 (Short ≤3, Medium 4–7, Long >7 đêm)"),
])
blank()
insert_image("Task 8.2 - Heatmap Occupancy.png",
             "Hình 8.2. Heatmap tỷ lệ lấp đầy theo tháng và nhóm minimum nights")
blank()
eval_section(
    express_bullets=[
        "PosX (Month – O): trục thời gian theo thứ tự → đúng.",
        "PosY (Minimum Nights Group – C): phân loại chính sách → đúng. Spatial region cho categorical.",
        "Color (Luminance/Sequential): thể hiện độ lớn occupancy (Q) → phù hợp. Sequential colormap (nhạt→đậm) đúng cho quantitative attribute có hướng sequential.",
    ],
    conclude_express="Kết luận: Heatmap là idiom phù hợp nhất khi cần phân tích 3 biến (2 keys + 1 value) trong cấu trúc ma trận.",
    effect_bullets=[
        "Accuracy: Color (Luminance) — accuracy thấp hơn Position, khó ước lượng chính xác % chênh lệch. Tuy nhiên mục tiêu là phát hiện pattern (cao/thấp), không phải đọc giá trị chính xác — trade-off chấp nhận được.",
        "Discriminability: 36 ô vuông đều nhau. Sequential palette giúp phân biệt ~5–7 cấp độ màu.",
        "Separability: PosX, PosY và Color kết hợp tốt. Area đồng đều không tạo bias về kích thước.",
    ],
    conclude_effect="",
    phanTich_bullets=[
        "Short minimum stay (≤3 đêm) có occupancy cao và ổn định quanh năm — chính sách linh hoạt thu hút nhiều khách nhất, đặc biệt trong mùa thấp điểm.",
        "Long minimum stay (>7 đêm) có occupancy thấp hơn trong phần lớn các tháng, nhưng tăng vào cuối năm — phù hợp cho khách công tác dài hạn.",
        "Tháng 11–12 có màu đậm nhất ở hầu hết các nhóm — xác nhận high season với nhu cầu cao bất kể chính sách minimum nights.",
        "Khuyến nghị cho host: Tháng 1–4 nên chuyển sang Short minimum stay; tháng 11–12 có thể giữ Medium/Long stay vì cầu cao tự nhiên bù đắp tính hạn chế của chính sách.",
    ]
)

# save
out_path = r"C:\Users\Duy\Desktop\dvfinal\docs\report\domain_task_5678.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
